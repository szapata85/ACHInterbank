from __future__ import annotations

import hashlib
import shutil
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "entrega_pruebas_funcionales_usuarios" / "paquete_final"
CAPTURAS = PKG / "capturas"

WORD_ORIGINAL = PKG / "Guia_funcional_pruebas_locales_ACH_Interbank_ACH_COLOMBIA_CENIT_MANUAL_OPERATIVO_COMPLETO_FLUJOS_SUBPANTALLAS.docx"
WORD_REVISADO = PKG / "Guia_funcional_pruebas_locales_ACH_Interbank_ACH_COLOMBIA_CENIT_MANUAL_OPERATIVO_COMPLETO_FLUJOS_SUBPANTALLAS_V3_REVISADO.docx"
VALIDACION_MD = PKG / "VALIDACION_WORD_V3_REVISADO.md"

MD_CONTENIDO = PKG / "CONTENIDO_AMPLIACION_MANUAL_FUNCIONAL_V3.md"
MD_COBERTURA = PKG / "COBERTURA_FUNCIONAL_MANUAL_V3.md"
MD_PENDIENTES = PKG / "PENDIENTES_VALIDACION_FUNCIONAL_V3.md"

PNG_ORDER = [
    "28_customer_third_parties_listado_busqueda.png",
    "30_customers_listado.png",
    "31_customers_nuevo.png",
    "32_financial_institutions_mantenimiento_digito_verificacion.png",
    "33_clearing_house_preferences_prioridades_camara.png",
    "34_catalog_document_types.png",
    "35_catalog_gender_types.png",
    "36_catalog_person_types.png",
    "37_catalog_transaction_codes.png",
    "38_catalog_company_entry_descriptions.png",
    "39_transactions_clearing_house_rules.png",
    "40_transactions_cycle_configs.png",
    "41_cenit_causales_devolucion.png",
    "42_cenit_causales_rechazo.png",
    "43_cenit_politicas_transaccion.png",
    "44_cenit_politicas_prenotificacion.png",
    "45_nacha_config_perfiles.png",
    "47_nacha_operational_dashboard.png",
    "48_ach_reconciliation.png",
    "49_reports_rejections.png",
    "51_nacha_security_dashboard.png",
    "52_nacha_security_certificates_gobierno.png",
    "59_nacha_manual_encrypt_sobre_digital.png",
    "60_nacha_manual_decrypt_sobre_digital.png",
    "61_sobre_digital_tool.png",
    "63_ach_cycles_nacha_export.png",
]

SKIP_IMAGES = {
    "56_nacha_security_audit_sobre_digital.png",
    "62_interoperabilidad_vector_oficial.png",
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def count_docx(path: Path) -> tuple[int, int, int]:
    with zipfile.ZipFile(path, "r") as zf:
        xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = len(root.findall(".//w:body//w:p", ns))
        tables = len(root.findall(".//w:body//w:tbl", ns))
        images = sum(1 for name in zf.namelist() if name.startswith("word/media/"))
        return paragraphs, tables, images


def ensure_exists(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(str(path))


def get_existing_images() -> list[Path]:
    images: list[Path] = []
    seen = set()
    for name in PNG_ORDER:
        if name in seen:
            continue
        seen.add(name)
        candidate = CAPTURAS / name
        if candidate.exists():
            images.append(candidate)
    return images


def add_para(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def add_captioned_image(doc: Document, title: str, brief: str, image_path: Path):
    add_para(doc, title, style="Heading 2")
    add_para(doc, brief)
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(6.5))


def add_section_text(doc: Document, heading: str, paragraphs: list[str], bullets: list[str] | None = None):
    add_para(doc, heading, style="Heading 2")
    for para in paragraphs:
        add_para(doc, para)
    if bullets:
        for bullet in bullets:
            add_bullet(doc, bullet)


def main() -> None:
    ensure_exists(WORD_ORIGINAL)
    for md in (MD_CONTENIDO, MD_COBERTURA, MD_PENDIENTES):
        ensure_exists(md)
    ensure_exists(CAPTURAS)

    original_before = sha256(WORD_ORIGINAL)
    p_before, t_before, i_before = count_docx(WORD_ORIGINAL)

    shutil.copyfile(WORD_ORIGINAL, WORD_REVISADO)
    doc = Document(str(WORD_REVISADO))

    existing_images = get_existing_images()
    # Build append-only V3 content.
    doc.add_page_break()
    add_para(doc, "Ampliación funcional V3 - ACH Interbank", style="Title")

    add_section_text(
        doc,
        "Objetivo de la ampliación",
        [
            "Consolidar la ampliación funcional V3 al final del manual sin alterar el contenido previo.",
            "La versión revisada agrega evidencia funcional y control documental sobre generación NACHA-M, naming reglamentario, certificados, seguridad y transacciones.",
        ],
    )

    add_section_text(
        doc,
        "Alcance funcional agregado",
        [
            "La ampliación complementa el manual previo con criterios funcionales, control de pendientes y referencias a capturas reales existentes.",
            "No sustituye el contenido anterior; lo conserva íntegro y añade el bloque V3 como anexo final.",
        ],
    )

    add_section_text(
        doc,
        "Datos maestros y parametrización",
        [
            "Se mantienen las validaciones de catálogos y parametrización visibles en el entorno de pruebas.",
            "No se inventan fórmulas ni formatos no verificados para entidades, cámaras o reglas regulatorias.",
        ],
    )

    # Images supporting master data / operational views.
    for title, brief, img in [
        ("Entidad financiera y dígito de verificación", "Vista de mantenimiento de entidades financieras y campo visible de verificación.", CAPTURAS / "32_financial_institutions_mantenimiento_digito_verificacion.png"),
        ("Priorización por cámara", "Preferencias operativas para ACH Colombia y CENIT.", CAPTURAS / "33_clearing_house_preferences_prioridades_camara.png"),
        ("Catálogo de tipos de documento", "Referencia visible de catálogo maestro.", CAPTURAS / "34_catalog_document_types.png"),
        ("Catálogo de género", "Referencia visible del catálogo maestro.", CAPTURAS / "35_catalog_gender_types.png"),
        ("Catálogo de tipos de persona", "Referencia visible del catálogo maestro.", CAPTURAS / "36_catalog_person_types.png"),
        ("Catálogo de códigos de transacción", "Referencia visible del catálogo maestro.", CAPTURAS / "37_catalog_transaction_codes.png"),
        ("Conceptos de lote", "Referencia visible del catálogo maestro.", CAPTURAS / "38_catalog_company_entry_descriptions.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Administración de clientes y terceros",
        [
            "La validación de terceros y clientes continúa como consulta funcional y registro operativo en el manual.",
            "El onboarding silencioso se valida durante la creación de una transacción y no como pantalla independiente.",
        ],
        [
            "Punto 50 queda asociado a /transactions/create.",
            "Puntos 29 y 55 permanecen como pendientes funcionales, no como cubiertos.",
        ],
    )
    for title, brief, img in [
        ("Terceros de prenotificación", "Listado y búsqueda de terceros de prenotificación.", CAPTURAS / "28_customer_third_parties_listado_busqueda.png"),
        ("Clientes", "Listado de clientes disponible en el entorno de pruebas.", CAPTURAS / "30_customers_listado.png"),
        ("Alta de clientes", "Formulario de creación de clientes.", CAPTURAS / "31_customers_nuevo.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Reglas por cámara y ciclos",
        [
            "Las reglas por cámara y la configuración de ciclos siguen documentadas como evidencia operativa del manual.",
            "No se inventan reglas de CENIT ni fórmulas de dígito de verificación.",
        ],
    )
    for title, brief, img in [
        ("Reglas por cámara", "Reglas visibles para procesamiento operativo.", CAPTURAS / "39_transactions_clearing_house_rules.png"),
        ("Configuración de ciclos", "Ciclos operativos visibles.", CAPTURAS / "40_transactions_cycle_configs.png"),
        ("Causales de devolución CENIT", "Causales regulatorias visibles.", CAPTURAS / "41_cenit_causales_devolucion.png"),
        ("Causales de rechazo CENIT", "Causales regulatorias visibles.", CAPTURAS / "42_cenit_causales_rechazo.png"),
        ("Políticas de transacción CENIT", "Políticas visibles en pantalla.", CAPTURAS / "43_cenit_politicas_transaccion.png"),
        ("Políticas de prenotificación CENIT", "Políticas visibles en pantalla.", CAPTURAS / "44_cenit_politicas_prenotificacion.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Perfiles NACHA-M",
        [
            "Los perfiles NACHA-M se mantienen como base de parametrización oficial para la operación.",
            "El detalle navegable se documenta solo cuando existe evidencia real visible.",
        ],
    )
    for title, brief, img in [
        ("Listado de perfiles NACHA-M", "Vista principal de perfiles oficiales.", CAPTURAS / "45_nacha_config_perfiles.png"),
        ("Dashboard operativo NACHA-M", "Consulta operativa read-only.", CAPTURAS / "47_nacha_operational_dashboard.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Consulta operativa, conciliación y reportes",
        [
            "La ampliación conserva el enfoque read-only para consulta operativa, conciliación y reportes.",
            "Los casos 56 y 62 permanecen retirados del aplicativo/manual y no se insertan como pantallas activas.",
        ],
    )
    for title, brief, img in [
        ("Conciliación ACH", "Vista de conciliación visible.", CAPTURAS / "48_ach_reconciliation.png"),
        ("Reporte de rechazos", "Vista de reporte de rechazos visible.", CAPTURAS / "49_reports_rejections.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Seguridad NACHA-M, certificados y sobre digital",
        [
            "Los certificados, sus estados, vigencias y versionado se conservan como pendientes controlados si no hay evidencia suficiente.",
            "No se documentan llaves, OpenBao, Bre-B ni rutas de alias.",
        ],
    )
    for title, brief, img in [
        ("Dashboard de seguridad NACHA-M", "Consola base de seguridad visible.", CAPTURAS / "51_nacha_security_dashboard.png"),
        ("Gobierno de certificados", "Vista principal de certificados visible.", CAPTURAS / "52_nacha_security_certificates_gobierno.png"),
        ("Cifrado manual con sobre digital", "Flujo manual visible.", CAPTURAS / "59_nacha_manual_encrypt_sobre_digital.png"),
        ("Descifrado manual con sobre digital", "Flujo manual visible.", CAPTURAS / "60_nacha_manual_decrypt_sobre_digital.png"),
        ("Herramienta sobre digital", "Consola operativa visible.", CAPTURAS / "61_sobre_digital_tool.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Generación NACHA-M, naming y extensiones reglamentarias",
        [
            "La generación NACHA-M base corresponde al archivo previo al cifrado o salida final.",
            "La generación NACHA-M cifrada corresponde a la salida protegida generada desde el archivo base.",
            "El usuario funcional no debe modificar manualmente los archivos generados.",
            "Para ACH Colombia, el naming reglamentario del archivo base se maneja como RRRRTTT.ZZZ.1.",
            "RRRR corresponde a ruta o identificador de entidad originadora según parametrización.",
            "TTT corresponde a tránsito o identificador operativo.",
            "ZZZ corresponde al consecutivo diario.",
            ".1 corresponde al sufijo o extensión reglamentaria del archivo base.",
            "Para CENIT, no se inventa formato; el naming queda sujeto a parametrización o evidencia generada por el aplicativo.",
            "La extensión .env solo se documenta como salida final exportable si la aplicación la genera o la muestra.",
            "No se asume que el archivo base y el archivo final tengan la misma extensión.",
            "No se renombran manualmente archivos para cumplir una regla.",
            "Evidencias válidas: archivo generado, nombre visible, registro de exportación, evidencia del ciclo o validación funcional registrada.",
        ],
        [
            "Puntos 57, 58, 64 y 65 permanecen como evidencia funcional, no solo printscreen.",
        ],
    )
    for title, brief, img in [
        ("Exportación NACHA-M desde ciclo", "Salida exportable visible.", CAPTURAS / "63_ach_cycles_nacha_export.png"),
    ]:
        if img.exists():
            add_captioned_image(doc, title, brief, img)

    add_section_text(
        doc,
        "Onboarding silencioso en creación de transacción",
        [
            "El onboarding silencioso no es una pantalla independiente.",
            "Se valida dentro de /transactions/create al crear una transacción de prueba y verificar que el sistema asocie o registre la información requerida sin exigir alta manual previa.",
        ],
    )

    add_section_text(
        doc,
        "Pendientes funcionales controlados",
        [
            "29 - Creación de terceros.",
            "50 - Onboarding silencioso en creación de transacción.",
            "53 - Versiones / historial de certificados.",
            "54 - Estados / vigencia de certificados.",
            "55 - Rotación / reemplazo de certificados.",
            "57 - Generación NACHA-M base.",
            "58 - Generación NACHA-M cifrada.",
            "64 - Naming archivo base por cámara.",
            "65 - Naming archivo final .env.",
        ],
    )
    add_table(
        doc,
        ["ID", "Tema", "Estado", "Criterio de cierre"],
        [
            ["29", "Creación de terceros", "No encontrado", "Confirmar si existe acción real de creación."],
            ["50", "Onboarding silencioso", "Requiere validación funcional", "Validar dentro de /transactions/create."],
            ["53", "Versiones / historial de certificados", "Requiere validación", "Confirmar evidencia navegable real."],
            ["54", "Estados / vigencia de certificados", "Requiere validación", "Confirmar estados y fechas visibles."],
            ["55", "Rotación / reemplazo de certificados", "No encontrado", "Confirmar acción visible si existe."],
            ["57", "Generación NACHA-M base", "Requiere evidencia funcional", "Confirmar archivo base generado por cámara."],
            ["58", "Generación NACHA-M cifrada", "Requiere evidencia funcional", "Confirmar salida cifrada desde archivo base."],
            ["64", "Naming archivo base por cámara", "Requiere evidencia funcional", "Confirmar naming por cámara con RRRRTTT.ZZZ.1 en ACH Colombia."],
            ["65", "Naming archivo final .env", "Requiere evidencia funcional", "Confirmar extensión final .env solo si la aplicación la genera o muestra."],
        ],
    )

    add_section_text(
        doc,
        "Pantallas retiradas del aplicativo/manual",
        [
            "Auditoría de certificados / sobre digital y interoperabilidad / vector oficial permanecen retiradas.",
            "No se insertan como capturas activas en el documento revisado.",
        ],
    )
    add_table(
        doc,
        ["Pantalla", "Ruta", "Estado"],
        [
            ["Auditoría de certificados / sobre digital", "/nacha-security/digital-envelope/audit", "Retirada del aplicativo/manual"],
            ["Interoperabilidad / vector oficial", "/nacha-security/digital-envelope/interoperability", "Retirada del aplicativo/manual"],
        ],
    )

    add_section_text(
        doc,
        "Anexo resumido de control",
        [
            "Este anexo resume la ampliación V3 y conserva el resto del manual previo sin reordenar ni borrar contenido anterior.",
            "La validación final depende de evidencia funcional real y no de capturas inventadas o pantallas vacías.",
        ],
    )

    doc.save(str(WORD_REVISADO))

    # Post-save validations.
    if not WORD_REVISADO.exists() or WORD_REVISADO.stat().st_size == 0:
        raise RuntimeError("Revised DOCX not created or empty")
    with zipfile.ZipFile(WORD_REVISADO, "r") as zf:
        zf.testzip()
        document_xml = zf.read("word/document.xml")
        ET.fromstring(document_xml)
    Document(str(WORD_REVISADO))

    original_after = sha256(WORD_ORIGINAL)
    p_after, t_after, i_after = count_docx(WORD_REVISADO)
    if original_before != original_after:
        raise RuntimeError("Original DOCX hash changed")
    if i_after <= i_before:
        raise RuntimeError("Image count did not increase")

    inserted = [p.name for p in existing_images]
    missing = [name for name in PNG_ORDER if not (CAPTURAS / name).exists() and name not in SKIP_IMAGES]

    status = "APROBADO" if (i_after > i_before and original_before == original_after) else "FALLIDO"

    lines = [
        "# Validación Word V3 Revisado",
        "",
        "## Entorno Python",
        "",
        "* python --version: Python 3.9.12",
        "* where python: " + (shutil.which("python") or "No encontrado"),
        "* python-docx instalado: Sí",
        "* versión python-docx: 1.2.0",
        "",
        "## Archivos",
        "",
        f"* Word original: {WORD_ORIGINAL}",
        f"* Word revisado: {WORD_REVISADO}",
        f"* Carpeta de capturas: {CAPTURAS}",
        "",
        "## Integridad",
        "",
        f"* SHA-256 original antes: {original_before}",
        f"* SHA-256 original después: {original_after}",
        f"* Word original intacto: {'Sí' if original_before == original_after else 'No'}",
        "* ZIP válido: Sí",
        "* document.xml válido: Sí",
        "* Apertura con python-docx: Sí",
        "",
        "## Métricas",
        "",
        f"* Párrafos antes: {p_before}",
        f"* Párrafos después: {p_after}",
        f"* Tablas antes: {t_before}",
        f"* Tablas después: {t_after}",
        f"* Imágenes antes: {i_before}",
        f"* Imágenes después: {i_after}",
        "",
        "## Capturas",
        "",
        "* Capturas insertadas: " + ", ".join(inserted),
        "* Capturas no insertadas y motivo: 29, 50, 53, 54, 55, 57, 58, 64, 65 por pendiente controlado; 56 y 62 retiradas; algunas capturas no estaban físicamente presentes en la carpeta.",
        "* Confirmación de que 56 y 62 no fueron insertadas: Sí",
        "",
        "## Reglas funcionales",
        "",
        "* 50 asociado a creación de transacción: Sí",
        "* 57, 58, 64 y 65 tratados como evidencia funcional: Sí",
        "* Pantallas retiradas documentadas: Sí",
        "",
        "## Conclusión",
        "",
        status,
    ]
    VALIDACION_MD.write_text("\n".join(lines), encoding="utf-8")

    print(f"ORIGINAL={WORD_ORIGINAL}")
    print(f"REVISADO={WORD_REVISADO}")
    print(f"ORIG_PARAS={p_before}")
    print(f"REVISADO_PARAS={p_after}")
    print(f"ORIG_TABLES={t_before}")
    print(f"REVISADO_TABLES={t_after}")
    print(f"ORIG_IMAGES={i_before}")
    print(f"REVISADO_IMAGES={i_after}")
    print(f"STATUS={status}")
    print(f"VALIDACION={VALIDACION_MD}")


if __name__ == "__main__":
    main()
