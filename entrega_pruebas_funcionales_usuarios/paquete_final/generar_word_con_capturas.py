from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
BASE_DOCX = BASE_DIR / "Guia_funcional_pruebas_locales_ACH_Interbank_CENIT.docx"
OUTPUT_DOCX = BASE_DIR / "Guia_funcional_pruebas_locales_ACH_Interbank_CENIT_FINAL_CON_CAPTURAS.docx"
CAPTURES_DIR = BASE_DIR / "capturas"

CAPTURES = [
    (
        "01_inicio_o_login.png",
        "Pantalla de ingreso",
        "Pantalla inicial para acceder al sistema con usuario de prueba.",
    ),
    (
        "02_dashboard.png",
        "Panel principal",
        "Pantalla principal despues del ingreso al sistema.",
    ),
    (
        "03_dashboard_operacional_nacha.png",
        "Dashboard operacional NACHA-M",
        "Vista de seguimiento funcional de informacion operacional NACHA-M.",
    ),
    (
        "04_configuracion_perfiles_nacha.png",
        "Configuracion de perfiles NACHA-M",
        "Pantalla para revisar la parametrizacion funcional de perfiles.",
    ),
    (
        "05_exportacion_nacha.png",
        "Exportacion NACHA",
        "Pantalla de apoyo para validaciones relacionadas con exportacion.",
    ),
    (
        "06_ciclos_ach.png",
        "Ciclos ACH",
        "Pantalla para revisar ciclos disponibles desde el menu del sistema.",
    ),
    (
        "07_transacciones.png",
        "Transacciones",
        "Pantalla para consultar y validar transacciones desde el menu del sistema.",
    ),
    (
        "08_cenit.png",
        "CENIT",
        "Pantalla de apoyo para pruebas funcionales relacionadas con CENIT.",
    ),
    (
        "09_uat_console.png",
        "Consola UAT",
        "Pantalla de apoyo para validaciones funcionales controladas.",
    ),
    (
        "10_menu_o_navegacion.png",
        "Menu y navegacion",
        "Referencia visual de las opciones principales disponibles.",
    ),
    (
        "11_uat_inbound_simulator.png",
        "Simulador de entrada UAT",
        "Pantalla de apoyo para pruebas funcionales de entrada controlada.",
    ),
]


def set_run_font(run, name="Arial", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clear_document(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for style_name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("Guia funcional de pruebas locales")
    set_run_font(r, size=24, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ACH Interbank / CENIT")
    set_run_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("Entrega para usuarios funcionales")
    set_run_font(r, size=13)


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True)


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"- {text}")
    set_run_font(r, size=11)


def add_capture_page(doc, image_path, title, description):
    doc.add_page_break()

    title_p = doc.add_paragraph(style="Heading 1")
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=16, bold=True)

    desc_p = doc.add_paragraph(style="Normal")
    desc_p.paragraph_format.space_after = Pt(10)
    desc_p.paragraph_format.line_spacing = 1.15
    desc_run = desc_p.add_run(description)
    set_run_font(desc_run, size=11)

    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_after = Pt(6)
    image_p.add_run().add_picture(str(image_path), width=Inches(6.1))


def main():
    if not BASE_DOCX.exists():
        raise FileNotFoundError(f"No existe el archivo base: {BASE_DOCX}")

    missing = [name for name, _, _ in CAPTURES if not (CAPTURES_DIR / name).exists()]
    if missing:
        raise FileNotFoundError(f"Faltan capturas requeridas: {', '.join(missing)}")

    doc = Document(str(BASE_DOCX))
    clear_document(doc)
    style_document(doc)

    add_title_page(doc)
    add_heading(doc, "Objetivo")
    add_body(
        doc,
        "Este documento orienta la ejecucion de pruebas locales funcionales y sirve como apoyo visual para el recorrido de las pantallas principales.",
    )

    add_heading(doc, "Ingreso al sistema")
    add_bullet(doc, "Ingresar por http://localhost:743/login.")
    add_bullet(doc, "Usar credenciales de prueba autorizadas.")
    add_bullet(doc, "No usar datos reales sensibles.")

    add_heading(doc, "Uso de la guia")
    add_bullet(doc, "Leer la guia antes de iniciar las pruebas.")
    add_bullet(doc, "Ejecutar los escenarios definidos en la matriz.")
    add_bullet(doc, "Registrar hallazgos en el formato de incidencias.")
    add_bullet(doc, "Adjuntar evidencia cuando ayude a explicar el caso.")
    add_bullet(doc, "Abrir Ciclos ACH y Transacciones desde el menu del sistema.")
    add_bullet(doc, "La ruta /auth/login no es una ruta para usuarios.")

    add_heading(doc, "Capturas de referencia")
    add_body(
        doc,
        "Las siguientes imagenes sirven como referencia visual para ubicar las pantallas principales incluidas en esta entrega.",
    )

    for filename, title, description in CAPTURES:
        add_capture_page(doc, CAPTURES_DIR / filename, title, description)

    doc.add_page_break()
    add_heading(doc, "Recomendaciones para usuarios")
    add_bullet(doc, "Probar solo los escenarios indicados en la matriz.")
    add_bullet(doc, "Registrar observaciones claras y completas.")
    add_bullet(doc, "No modificar configuraciones no solicitadas.")
    add_bullet(doc, "No usar informacion real sensible.")
    add_bullet(doc, "Reportar errores con captura y descripcion.")

    add_heading(doc, "Cierre")
    add_body(
        doc,
        "El paquete queda listo para iniciar pruebas locales funcionales con apoyo documental y visual.",
    )

    doc.save(str(OUTPUT_DOCX))
    print(f"Archivo generado: {OUTPUT_DOCX}")
    print(f"Capturas incrustadas: {len(CAPTURES)}")


if __name__ == "__main__":
    main()
