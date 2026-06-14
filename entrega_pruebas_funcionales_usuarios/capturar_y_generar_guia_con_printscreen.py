#!/usr/bin/env python3
"""
Captura pantallas funcionales de ACH Interbank con Chromium/Playwright y genera
una copia de la guía Word con las imágenes anexadas.

Uso sugerido:
  python capturar_y_generar_guia_con_printscreen.py \
    --base-url http://localhost:743 \
    --username admin \
    --password "CONTRASENA_DE_PRUEBA" \
    --docx-in Guia_funcional_pruebas_locales_ACH_Interbank_CENIT.docx \
    --docx-out Guia_funcional_pruebas_locales_ACH_Interbank_CENIT_con_capturas_reales.docx

También puede usar la variable de entorno ACH_TEST_PASSWORD para no escribir la clave en consola:
  Windows PowerShell:  $env:ACH_TEST_PASSWORD="CONTRASENA_DE_PRUEBA"
  Linux/macOS:         export ACH_TEST_PASSWORD="CONTRASENA_DE_PRUEBA"
"""

from __future__ import annotations

import argparse
import os
import re
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
except Exception as exc:  # pragma: no cover
    print("ERROR: Playwright no está instalado.")
    print("Instale dependencias con: pip install playwright python-docx")
    print("Luego ejecute: python -m playwright install chromium")
    raise exc


@dataclass
class ScreenSpec:
    code: str
    title: str
    route: str
    expected: str
    needs_login: bool = True


DEFAULT_SCREENS: list[ScreenSpec] = [
    ScreenSpec("01", "Pantalla inicial / Login", "/", "Debe verse la pantalla inicial del sistema o formulario de ingreso.", False),
    ScreenSpec("02", "Menú principal", "/", "Debe verse el menú principal y las opciones habilitadas para el usuario funcional.", True),
    ScreenSpec("03", "Cámaras de compensación", "/clearing-houses", "Deben visualizarse ACH Colombia y CENIT de forma clara.", True),
    ScreenSpec("04", "Ciclos ACH / CENIT", "/ach-cycles", "Deben visualizarse ciclos, horarios y estado funcional.", True),
    ScreenSpec("05", "Configuración funcional de ciclos", "/transactions/cycle-configs", "Deben verse configuraciones de ciclos y acciones disponibles, si aplica.", True),
    ScreenSpec("06", "Transacciones", "/transactions", "Debe cargar la grilla de transacciones con filtros, columnas y acciones.", True),
    ScreenSpec("07", "Entidades financieras", "/financial-institutions", "Deben visualizarse entidades financieras de prueba y CFA como entidad destino cuando aplique.", True),
    ScreenSpec("08", "Descripciones / conceptos de entrada", "/company-entry-descriptions", "Deben visualizarse conceptos o descripciones disponibles para prueba.", True),
]


def normalize_base_url(base_url: str) -> str:
    base_url = base_url.strip()
    if not base_url.startswith(("http://", "https://")):
        base_url = "http://" + base_url
    if not base_url.endswith("/"):
        base_url += "/"
    return base_url


def build_url(base_url: str, route: str) -> str:
    return urljoin(base_url, route.lstrip("/"))


def slugify(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^a-z0-9áéíóúñü]+", "_", value, flags=re.IGNORECASE)
    value = value.strip("_")
    return value[:80] or "captura"


def first_visible(page, selectors: Iterable[str], timeout_ms: int = 1000):
    for selector in selectors:
        try:
            loc = page.locator(selector).first
            if loc.count() > 0:
                loc.wait_for(state="visible", timeout=timeout_ms)
                return loc
        except Exception:
            continue
    return None


def try_login(page, username: str | None, password: str | None) -> bool:
    if not username or not password:
        return False

    user_selectors = [
        'input[name*="user" i]',
        'input[id*="user" i]',
        'input[placeholder*="usuario" i]',
        'input[placeholder*="user" i]',
        'input[type="text"]',
        'input[type="email"]',
    ]
    password_selectors = [
        'input[type="password"]',
        'input[name*="password" i]',
        'input[id*="password" i]',
        'input[placeholder*="contraseña" i]',
        'input[placeholder*="password" i]',
    ]

    user_input = first_visible(page, user_selectors)
    pass_input = first_visible(page, password_selectors)
    if not user_input or not pass_input:
        return False

    user_input.fill(username)
    pass_input.fill(password)

    submit_selectors = [
        'button[type="submit"]',
        'button:has-text("Ingresar")',
        'button:has-text("Iniciar")',
        'button:has-text("Entrar")',
        'button:has-text("Login")',
        'input[type="submit"]',
    ]
    button = first_visible(page, submit_selectors)
    if not button:
        try:
            pass_input.press("Enter")
        except Exception:
            return False
    else:
        button.click()

    try:
        page.wait_for_load_state("networkidle", timeout=15000)
    except PlaywrightTimeoutError:
        pass
    time.sleep(1)
    return True


def capture_screens(args) -> list[tuple[ScreenSpec, Path, str]]:
    base_url = normalize_base_url(args.base_url)
    out_dir = Path(args.screens_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    password = args.password or os.environ.get("ACH_TEST_PASSWORD")
    captures: list[tuple[ScreenSpec, Path, str]] = []

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=not args.no_headless)
        context = browser.new_context(viewport={"width": args.width, "height": args.height})
        page = context.new_page()

        # 1. Pantalla inicial / login
        first = DEFAULT_SCREENS[0]
        first_path = out_dir / f"{first.code}_{slugify(first.title)}.png"
        try:
            page.goto(build_url(base_url, first.route), wait_until="networkidle", timeout=args.timeout)
        except PlaywrightTimeoutError:
            page.goto(build_url(base_url, first.route), wait_until="domcontentloaded", timeout=args.timeout)
        page.screenshot(path=str(first_path), full_page=True)
        captures.append((first, first_path, "OK"))

        logged = try_login(page, args.username, password)
        if args.username and password and not logged:
            print("ADVERTENCIA: No fue posible detectar automáticamente los campos de login. Se capturarán rutas, pero podrían mostrar la pantalla de ingreso.")
        elif not password:
            print("ADVERTENCIA: No se recibió contraseña. Solo se garantiza la captura de la pantalla inicial/login.")

        # 2+. Rutas funcionales
        for spec in DEFAULT_SCREENS[1:]:
            path = out_dir / f"{spec.code}_{slugify(spec.title)}.png"
            url = build_url(base_url, spec.route)
            status = "OK"
            try:
                page.goto(url, wait_until="networkidle", timeout=args.timeout)
            except PlaywrightTimeoutError:
                status = "OK con espera parcial"
                try:
                    page.goto(url, wait_until="domcontentloaded", timeout=args.timeout)
                except Exception as exc:
                    status = f"ERROR: {exc}"
            except Exception as exc:
                status = f"ERROR: {exc}"
            try:
                page.screenshot(path=str(path), full_page=True)
            except Exception as exc:
                status = f"ERROR al capturar: {exc}"
            captures.append((spec, path, status))

        context.close()
        browser.close()
    return captures


def set_normal_style(doc: Document):
    try:
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)
    except Exception:
        pass


def append_visual_annex(docx_in: Path, docx_out: Path, captures: list[tuple[ScreenSpec, Path, str]], base_url: str):
    doc = Document(str(docx_in))
    set_normal_style(doc)

    doc.add_page_break()
    title = doc.add_heading("Anexo visual - Printscreen de pruebas locales", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("Objetivo: ").bold = True
    p.add_run("dejar evidencia visual de las pantallas principales que deben revisar los usuarios funcionales durante las pruebas locales.")

    p = doc.add_paragraph()
    p.add_run("URL base usada para las capturas: ").bold = True
    p.add_run(normalize_base_url(base_url))

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Pantalla"
    hdr[2].text = "Ruta"
    hdr[3].text = "Resultado esperado"
    for spec, _path, status in captures:
        row = table.add_row().cells
        row[0].text = spec.code
        row[1].text = f"{spec.title}\nEstado captura: {status}"
        row[2].text = spec.route
        row[3].text = spec.expected

    for spec, path, status in captures:
        doc.add_heading(f"{spec.code}. {spec.title}", level=2)
        p = doc.add_paragraph()
        p.add_run("Ruta: ").bold = True
        p.add_run(spec.route)
        p.add_run(" | Resultado esperado: ").bold = True
        p.add_run(spec.expected)
        p.add_run(" | Estado de captura: ").bold = True
        p.add_run(status)

        if path.exists() and path.stat().st_size > 0:
            try:
                doc.add_picture(str(path), width=Inches(6.4))
            except Exception as exc:
                doc.add_paragraph(f"No fue posible insertar la imagen {path.name}: {exc}")
        else:
            doc.add_paragraph(f"Captura no disponible: {path.name}")

    docx_out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_out))


def main() -> int:
    parser = argparse.ArgumentParser(description="Captura printscreen ACH Interbank con Playwright y genera guía Word con imágenes.")
    parser.add_argument("--base-url", default="http://localhost:743", help="URL local de la SPA. Ej: http://localhost:743")
    parser.add_argument("--username", default="admin", help="Usuario de prueba autorizado. No usar usuarios productivos.")
    parser.add_argument("--password", default=None, help="Contraseña de prueba. Alternativa recomendada: variable ACH_TEST_PASSWORD.")
    parser.add_argument("--docx-in", default="Guia_funcional_pruebas_locales_ACH_Interbank_CENIT.docx", help="Guía Word base.")
    parser.add_argument("--docx-out", default="Guia_funcional_pruebas_locales_ACH_Interbank_CENIT_con_capturas_reales.docx", help="Guía Word de salida.")
    parser.add_argument("--screens-dir", default="capturas_ach_interbank", help="Carpeta donde se guardan los PNG.")
    parser.add_argument("--timeout", type=int, default=25000, help="Timeout Playwright en milisegundos.")
    parser.add_argument("--width", type=int, default=1440, help="Ancho del navegador para capturas.")
    parser.add_argument("--height", type=int, default=1000, help="Alto del navegador para capturas.")
    parser.add_argument("--no-headless", action="store_true", help="Muestra el navegador mientras captura.")
    args = parser.parse_args()

    docx_in = Path(args.docx_in).resolve()
    if not docx_in.exists():
        print(f"ERROR: No existe el archivo Word base: {docx_in}")
        return 2

    print(f"Capturando pantallas desde {normalize_base_url(args.base_url)} ...")
    captures = capture_screens(args)

    docx_out = Path(args.docx_out).resolve()
    append_visual_annex(docx_in, docx_out, captures, args.base_url)

    print("\nProceso finalizado.")
    print(f"Capturas PNG: {Path(args.screens_dir).resolve()}")
    print(f"Guía generada: {docx_out}")
    print("Revise visualmente el documento antes de enviarlo a usuarios.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
